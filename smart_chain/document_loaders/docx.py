import os
from ..documents import Document
from docx import Document as DocxDocument


class Docx2txtLoader:
    def __init__(self, file_path):
        self.file_path = file_path

    def load(self):
        doc = DocxDocument(self.file_path)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        content = "\n".join(paragraphs)
        metadata = {"source": self.file_path}
        return [Document(page_content=content, metadata=metadata)]
